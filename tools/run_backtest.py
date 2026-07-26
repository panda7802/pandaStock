from __future__ import annotations

import base64
import csv
import json
import math
import re
import struct
from dataclasses import dataclass
from datetime import datetime, timedelta, timezone
from pathlib import Path
from typing import Iterable

import numpy as np
import pandas as pd
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
RAW_DIR = ROOT / "raw_data"
OUTPUT_DIR = ROOT / "策略回测输出"
DATA_DIR = OUTPUT_DIR / "历史数据"
RESULT_DIR = OUTPUT_DIR / "回测中间结果"
EXCEL_DIR = OUTPUT_DIR / "单标的Excel"
SINGLE_CHART_DIR = OUTPUT_DIR / "单标的折线图"
CROSS_CHART_DIR = OUTPUT_DIR / "跨标的折线图"
REPORT_DIR = OUTPUT_DIR / "报告"

START_YEARS = [2000, 2001, 2007, 2008, 2010, 2018, 2022, 2023]
TICKERS = ["QQQ", "TQQQ", "SQQQ", "DIA", "UDOW", "SDOW", "SPY", "UPRO", "SPXU", "KO", "BRK.B"]
STRATEGIES = ["A", "B", "C"]
BASE_SOURCES = {
    "QQQ": "https://totalrealreturns.com/n/QQQ?start=1999-03-10&end=2026-07-24",
    "DIA": "https://totalrealreturns.com/n/DIA?start=1998-01-20&end=2026-07-24",
    "SPY": "https://totalrealreturns.com/n/SPY?start=1993-01-29&end=2026-07-24",
    "KO": "https://totalrealreturns.com/n/KO?start=1962-01-02&end=2026-07-24",
    "BRK.B": "https://totalrealreturns.com/n/BRK-B?start=1996-05-09&end=2026-07-24",
}
LEVERAGED = {
    "TQQQ": ("QQQ", 3.0),
    "SQQQ": ("QQQ", -3.0),
    "UDOW": ("DIA", 3.0),
    "SDOW": ("DIA", -3.0),
    "UPRO": ("SPY", 3.0),
    "SPXU": ("SPY", -3.0),
}

COLORS = {
    "A": "#2563EB",
    "B": "#F97316",
    "C": "#16A34A",
}
CROSS_COLORS = [
    "#2563EB", "#7C3AED", "#DB2777", "#0F766E", "#65A30D", "#CA8A04",
    "#EA580C", "#DC2626", "#475569", "#0891B2", "#9333EA",
]


def ensure_dirs() -> None:
    for directory in [
        DATA_DIR,
        RESULT_DIR,
        EXCEL_DIR,
        SINGLE_CHART_DIR,
        CROSS_CHART_DIR,
        REPORT_DIR,
    ]:
        directory.mkdir(parents=True, exist_ok=True)


def parse_total_return_html(path: Path) -> pd.DataFrame:
    text = path.read_text(encoding="utf-8", errors="replace")
    dates_match = re.search(r"sharedDatesColumnInput\s*=\s*\[([^\]]*)\]", text, re.S)
    if not dates_match:
        raise ValueError(f"未找到日期列：{path}")
    date_deltas = [int(v.strip()) for v in dates_match.group(1).split(",") if v.strip()]
    if not date_deltas:
        raise ValueError(f"日期列为空：{path}")

    values_match = re.search(
        r"Growth of \$10,000.*?window\.decodeFloat32Column\(\"([A-Za-z0-9+/=]+)\"\)",
        text,
        re.S,
    )
    if not values_match:
        raise ValueError(f"未找到总收益序列：{path}")
    raw = base64.b64decode(values_match.group(1))
    if len(raw) % 4:
        raise ValueError(f"Float32 数据长度异常：{path}")
    values = np.asarray(struct.unpack(f"<{len(raw) // 4}f", raw), dtype=float)

    day_values = np.cumsum(np.asarray(date_deltas, dtype=np.int64))
    dates = pd.to_datetime(day_values, unit="D", origin="unix", utc=True).tz_localize(None)
    if len(dates) != len(values):
        raise ValueError(f"日期与价格长度不一致：{path} {len(dates)} != {len(values)}")

    frame = pd.DataFrame({"Date": dates, "Close": values})
    frame = frame.replace([np.inf, -np.inf], np.nan).dropna()
    frame = frame[frame["Close"] > 0].drop_duplicates("Date").sort_values("Date").reset_index(drop=True)
    if frame.empty:
        raise ValueError(f"无有效数据：{path}")
    return frame


def build_leveraged(base: pd.DataFrame, leverage: float) -> pd.DataFrame:
    returns = base["Close"].pct_change().fillna(0.0).to_numpy(dtype=float)
    leveraged_returns = np.maximum(-0.999999, leverage * returns)
    index = np.empty(len(base), dtype=float)
    index[0] = 100.0
    for i in range(1, len(base)):
        index[i] = index[i - 1] * (1.0 + leveraged_returns[i])
    return pd.DataFrame({"Date": base["Date"].copy(), "Close": index})


def save_history(symbol: str, frame: pd.DataFrame, source: str, underlying: str, leverage: float) -> None:
    out = frame.copy()
    out["DailyReturn"] = out["Close"].pct_change()
    out["Underlying"] = underlying
    out["Leverage"] = leverage
    out["Source"] = source
    out.to_csv(DATA_DIR / f"{symbol}.csv", index=False, encoding="utf-8-sig", date_format="%Y-%m-%d")


@dataclass
class StrategyMetrics:
    final_return: float
    cagr: float
    max_drawdown: float
    volatility: float
    trades: int
    invested_pct: float

    def to_dict(self) -> dict:
        return {
            "final_return": float(self.final_return),
            "cagr": float(self.cagr),
            "max_drawdown": float(self.max_drawdown),
            "volatility": float(self.volatility),
            "trades": int(self.trades),
            "invested_pct": float(self.invested_pct),
        }


def calculate_metrics(wealth: np.ndarray, daily_strategy_returns: np.ndarray, positions: np.ndarray, dates: pd.Series) -> StrategyMetrics:
    elapsed_days = max(1, int((dates.iloc[-1] - dates.iloc[0]).days))
    years = elapsed_days / 365.25
    final_return = wealth[-1] - 1.0
    cagr = wealth[-1] ** (1.0 / years) - 1.0 if years > 0 and wealth[-1] > 0 else float("nan")
    peaks = np.maximum.accumulate(wealth)
    max_drawdown = float(np.min(wealth / peaks - 1.0))
    volatility = float(np.std(daily_strategy_returns[1:], ddof=1) * math.sqrt(252)) if len(wealth) > 2 else 0.0
    trades = int(np.sum(np.abs(np.diff(positions)) > 0))
    invested_pct = float(np.mean(positions))
    return StrategyMetrics(final_return, cagr, max_drawdown, volatility, trades, invested_pct)


def run_backtest(frame: pd.DataFrame, start_year: int) -> tuple[pd.DataFrame, dict[str, StrategyMetrics], pd.DataFrame]:
    work = frame.copy()
    work["MA5"] = work["Close"].rolling(5, min_periods=5).mean()
    target = pd.Timestamp(year=start_year, month=1, day=1)
    candidates = work.index[work["Date"] >= target]
    if len(candidates) == 0:
        raise ValueError(f"{start_year} 后无数据")
    start_idx = int(candidates[0])
    if start_idx < 4:
        raise ValueError(f"{start_year} 起点前缺少 5 日均线所需历史")

    sl = work.iloc[start_idx:].copy().reset_index(drop=True)
    n = len(sl)
    close = sl["Close"].to_numpy(dtype=float)
    ma = sl["MA5"].to_numpy(dtype=float)
    daily_returns = np.zeros(n, dtype=float)
    daily_returns[1:] = close[1:] / close[:-1] - 1.0

    pos_a = np.ones(n, dtype=float)
    pos_b = np.ones(n, dtype=float)
    for i in range(1, n):
        prev_close = close[i - 1]
        prev_ma = ma[i - 1]
        if np.isfinite(prev_ma):
            if prev_close > prev_ma:
                pos_a[i] = 1.0
                pos_b[i] = 0.0
            elif prev_close < prev_ma:
                pos_a[i] = 0.0
                pos_b[i] = 1.0
            else:
                pos_a[i] = pos_a[i - 1]
                pos_b[i] = pos_b[i - 1]
        else:
            pos_a[i] = pos_a[i - 1]
            pos_b[i] = pos_b[i - 1]

    strategy_returns = {
        "A": pos_a * daily_returns,
        "B": pos_b * daily_returns,
        "C": daily_returns.copy(),
    }
    positions = {"A": pos_a, "B": pos_b, "C": np.ones(n, dtype=float)}
    wealth = {}
    for strategy in STRATEGIES:
        w = np.ones(n, dtype=float)
        for i in range(1, n):
            w[i] = w[i - 1] * (1.0 + strategy_returns[strategy][i])
        wealth[strategy] = w

    daily = pd.DataFrame(
        {
            "Date": sl["Date"],
            "Close": close,
            "MA5": ma,
            "A_Position": pos_a,
            "B_Position": pos_b,
            "A_Wealth": wealth["A"],
            "B_Wealth": wealth["B"],
            "C_Wealth": wealth["C"],
        }
    )
    daily["Month"] = daily["Date"].dt.to_period("M")
    monthly = daily.groupby("Month", sort=True).first().reset_index(drop=True)
    monthly = monthly[["Date", "A_Wealth", "B_Wealth", "C_Wealth"]].copy()
    monthly.rename(columns={"A_Wealth": "A", "B_Wealth": "B", "C_Wealth": "C"}, inplace=True)
    for strategy in STRATEGIES:
        monthly[strategy] = monthly[strategy] - 1.0

    metrics = {
        strategy: calculate_metrics(
            wealth[strategy], strategy_returns[strategy], positions[strategy], sl["Date"]
        )
        for strategy in STRATEGIES
    }
    return monthly, metrics, daily


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    candidates = [
        Path(r"C:\Windows\Fonts\msyhbd.ttc" if bold else r"C:\Windows\Fonts\msyh.ttc"),
        Path(r"C:\Windows\Fonts\simhei.ttf"),
        Path(r"C:\Windows\Fonts\arial.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size=size)
    return ImageFont.load_default()


def fmt_return(value: float) -> str:
    pct = value * 100.0
    sign = "+" if pct > 0 else ""
    ap = abs(pct)
    if ap >= 100_000_000:
        return f"{sign}{pct / 100_000_000:.1f}亿%"
    if ap >= 10_000:
        return f"{sign}{pct / 10_000:.1f}万%"
    if ap >= 1_000:
        return f"{sign}{pct:,.0f}%"
    return f"{sign}{pct:.1f}%"


def nice_linear_ticks(vmin: float, vmax: float, count: int = 7) -> list[float]:
    if math.isclose(vmin, vmax):
        return [vmin]
    raw = (vmax - vmin) / max(1, count - 1)
    exponent = math.floor(math.log10(abs(raw))) if raw else 0
    fraction = raw / (10 ** exponent)
    nice_fraction = 1 if fraction <= 1 else 2 if fraction <= 2 else 5 if fraction <= 5 else 10
    step = nice_fraction * (10 ** exponent)
    start = math.floor(vmin / step) * step
    end = math.ceil(vmax / step) * step
    ticks = []
    value = start
    while value <= end + step * 0.1 and len(ticks) < 30:
        ticks.append(value)
        value += step
    return ticks


def plot_lines(
    dates: list[pd.Timestamp],
    series: dict[str, list[float]],
    colors: dict[str, str],
    title: str,
    subtitle: str,
    output: Path,
) -> None:
    width, height = 1800, 1000
    image = Image.new("RGB", (width, height), "#FFFFFF")
    draw = ImageDraw.Draw(image)
    title_font = font(34, bold=True)
    subtitle_font = font(21)
    axis_font = font(18)
    small_font = font(16)
    legend_font = font(17)
    draw.text((85, 38), title, fill="#0F172A", font=title_font)
    draw.text((85, 88), subtitle, fill="#64748B", font=subtitle_font)

    left, top, right, bottom = 150, 155, 1390, 865
    all_values = [float(v) * 100.0 for values in series.values() for v in values if np.isfinite(v)]
    vmin = min(all_values)
    vmax = max(all_values)
    linear = vmax <= 1000 and vmin >= -150 and (vmax - vmin) <= 1400
    if linear:
        pad = max(5.0, (vmax - vmin) * 0.08)
        lo, hi = vmin - pad, vmax + pad
        transform = lambda value: value
        ticks = nice_linear_ticks(lo, hi)
    else:
        scale = 100.0
        transform = lambda value: math.copysign(math.log1p(abs(value) / scale), value)
        lo = transform(vmin)
        hi = transform(vmax)
        candidate_ticks = [-100, -50, -10, 0, 10, 50, 100, 500]
        positive = 1000.0
        while positive <= max(1000.0, vmax * 1.1) and len(candidate_ticks) < 20:
            candidate_ticks.append(positive)
            positive *= 10.0
        ticks = [t for t in candidate_ticks if lo <= transform(t) <= hi]
        if 0 not in ticks and lo <= 0 <= hi:
            ticks.append(0)
        ticks = sorted(set(ticks))
        draw.text((left, height - 45), "注：纵轴使用对称对数刻度，以同时展示极大收益与接近 -100% 的序列。", fill="#64748B", font=small_font)

    transformed_values = [transform(v) for v in all_values]
    ylo = min(transformed_values)
    yhi = max(transformed_values)
    if math.isclose(ylo, yhi):
        ylo -= 1
        yhi += 1
    pad_y = (yhi - ylo) * 0.06
    ylo -= pad_y
    yhi += pad_y

    def xcoord(index: int) -> float:
        return left + (right - left) * index / max(1, len(dates) - 1)

    def ycoord(value_pct: float) -> float:
        tv = transform(value_pct)
        return bottom - (bottom - top) * (tv - ylo) / (yhi - ylo)

    draw.rectangle((left, top, right, bottom), outline="#CBD5E1", width=2)
    for tick in ticks:
        ty = ycoord(float(tick))
        if top <= ty <= bottom:
            draw.line((left, ty, right, ty), fill="#E2E8F0", width=1)
            label = fmt_return(float(tick) / 100.0)
            bbox = draw.textbbox((0, 0), label, font=axis_font)
            draw.text((left - 18 - (bbox[2] - bbox[0]), ty - 10), label, fill="#475569", font=axis_font)

    tick_count = min(8, len(dates))
    x_indices = sorted(set(int(round(i * (len(dates) - 1) / max(1, tick_count - 1))) for i in range(tick_count)))
    for idx in x_indices:
        x = xcoord(idx)
        draw.line((x, bottom, x, bottom + 8), fill="#94A3B8", width=1)
        label = dates[idx].strftime("%Y-%m")
        bbox = draw.textbbox((0, 0), label, font=axis_font)
        draw.text((x - (bbox[2] - bbox[0]) / 2, bottom + 15), label, fill="#475569", font=axis_font)

    for label, values in series.items():
        points = []
        for idx, value in enumerate(values):
            if np.isfinite(value):
                points.append((xcoord(idx), ycoord(float(value) * 100.0)))
        if len(points) >= 2:
            draw.line(points, fill=colors[label], width=4, joint="curve")

    legend_x, legend_y = 1435, 175
    draw.text((legend_x, legend_y - 42), "图例 / 期末收益", fill="#0F172A", font=font(20, bold=True))
    for idx, (label, values) in enumerate(series.items()):
        y = legend_y + idx * 54
        draw.line((legend_x, y + 12, legend_x + 34, y + 12), fill=colors[label], width=5)
        end_value = next((v for v in reversed(values) if np.isfinite(v)), float("nan"))
        legend_text = f"{label}  {fmt_return(float(end_value))}"
        draw.text((legend_x + 46, y), legend_text, fill="#1E293B", font=legend_font)

    image.save(output, format="PNG", optimize=True)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def configure_document_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    for style_name, size, color, before, after in [
        ("Title", 28, "17365D", 0, 12),
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "17365D", 8, 4),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    header_props = table.rows[0]._tr.get_or_add_trPr()
    repeat_header = OxmlElement("w:tblHeader")
    repeat_header.set(qn("w:val"), "true")
    header_props.append(repeat_header)
    header_cells = table.rows[0].cells
    for idx, text in enumerate(headers):
        header_cells[idx].text = text
        shade_cell(header_cells[idx], "F2F4F7")
        header_cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(header_cells[idx])
        for run in header_cells[idx].paragraphs[0].runs:
            run.font.bold = True
            run.font.size = Pt(9)
    for row in rows:
        cells = table.add_row().cells
        row_props = table.rows[-1]._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        cant_split.set(qn("w:val"), "true")
        row_props.append(cant_split)
        for idx, text in enumerate(row):
            cells[idx].text = str(text)
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cells[idx])
            for paragraph in cells[idx].paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(9)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    return table


def build_report(summary: pd.DataFrame, end_date: str, results: list[dict]) -> Path:
    doc = Document()
    configure_document_styles(doc)
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.text = "策略回测研究｜5 日均线择时与长期持有"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(100, 116, 139)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("第 ")
    add_page_number(footer)
    footer.add_run(" 页")
    for run in footer.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(100, 116, 139)

    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run("三种策略与不同标的的关系\n")
    sub = p.add_run("2000–2023 多起点、11 个标的、月度收益对比")
    sub.font.size = Pt(16)
    sub.font.color.rgb = RGBColor(71, 85, 105)
    doc.add_paragraph(f"数据截至：{end_date}｜生成日期：{datetime.now().date().isoformat()}")
    doc.add_paragraph(
        "研究范围包括 QQQ、DIA、SPY、KO、BRK.B，以及基于对应底层逐日总收益模拟的 3 倍做多与 3 倍反向序列。"
    )
    doc.add_paragraph("研究用途：比较规则与标的特征；不构成投资建议。")
    doc.add_page_break()

    doc.add_heading("执行摘要", level=1)
    best_counts = (
        summary.sort_values(["Ticker", "StartYear", "FinalReturn"], ascending=[True, True, False])
        .groupby(["Ticker", "StartYear"])
        .first()
        .reset_index()["Strategy"]
        .value_counts()
    )
    avg_by_strategy = summary.groupby("Strategy").agg(
        MedianCAGR=("CAGR", "median"),
        MedianFinal=("FinalReturn", "median"),
        MedianDrawdown=("MaxDrawdown", "median"),
    )
    best_strategy = avg_by_strategy["MedianCAGR"].idxmax()
    safest_strategy = avg_by_strategy["MedianDrawdown"].idxmax()
    doc.add_paragraph(
        f"在 88 组“标的×起点”比较中，按期末收益胜出次数分别为："
        f"A {int(best_counts.get('A', 0))} 次、B {int(best_counts.get('B', 0))} 次、C {int(best_counts.get('C', 0))} 次。"
        f"跨样本中位年化收益最高的是策略 {best_strategy}；中位最大回撤最小的是策略 {safest_strategy}。"
    )
    doc.add_paragraph(
        "核心结论是：策略效果不能脱离标的方向性。长期正向漂移的宽基指数和优质公司更有利于持有或顺势；"
        "长期存在结构性衰减的每日反向 3 倍产品，则更依赖择时，长期持有通常处于劣势。"
    )
    doc.add_paragraph(
        "A 与 B 是同一信号的两面：A 在价格位于 5 日均线上方时持有，偏短周期顺势；B 在均线下方时持有，"
        "偏短周期逆向暴露。由于信号在收盘后确认并于下一交易日生效，本报告没有使用未来数据。"
    )

    doc.add_heading("方法与口径", level=1)
    methods = [
        "起点：2000、2001、2007、2008、2010、2018、2022、2023 年 1 月 1 日；若休市，使用其后第一个交易日。",
        "价格：使用分红再投资后的逐日总收益序列；均线与收益均基于该序列。",
        "策略 A：起点满仓；前一交易日收盘价高于 5 日均线则下一日持有，否则空仓。",
        "策略 B：起点满仓；前一交易日收盘价低于 5 日均线则下一日持有，否则空仓。",
        "策略 C：从起点起一直持有。",
        "相等处理：价格等于均线时沿用上一日仓位。现金收益为 0；不计交易费、税费与滑点。",
        "三倍序列：每日收益按底层总收益的 +3 倍或 -3 倍模拟，并将单日最低收益限制为 -99.9999%。"
        "未计管理费、融资成本、掉期价差和跟踪误差。",
        "月度记录：每个自然月首个交易日的累计收益；图表与 Excel 均采用同一组月度观测。",
    ]
    for item in methods:
        doc.add_paragraph(item, style=None).style = doc.styles["Normal"]
        doc.paragraphs[-1].style = doc.styles["Normal"]
        doc.paragraphs[-1].text = "• " + item

    cross_heading = doc.add_heading("跨标的汇总", level=1)
    cross_heading.paragraph_format.page_break_before = True
    ticker_summary = (
        summary.groupby(["Ticker", "Strategy"])
        .agg(
            MedianCAGR=("CAGR", "median"),
            MedianFinal=("FinalReturn", "median"),
            MedianDrawdown=("MaxDrawdown", "median"),
        )
        .reset_index()
    )
    rows = []
    for ticker in TICKERS:
        subdf = ticker_summary[ticker_summary["Ticker"] == ticker].copy()
        best = subdf.loc[subdf["MedianCAGR"].idxmax()]
        safest = subdf.loc[subdf["MedianDrawdown"].idxmax()]
        hold = subdf[subdf["Strategy"] == "C"].iloc[0]
        rows.append(
            [
                ticker,
                str(best["Strategy"]),
                f"{best['MedianCAGR']:.1%}",
                str(safest["Strategy"]),
                f"{hold['MedianFinal']:.1%}",
            ]
        )
    add_table(
        doc,
        ["标的", "中位年化最佳", "最佳中位 CAGR", "回撤最小策略", "C 策略中位期末收益"],
        rows,
        [0.8, 1.15, 1.25, 1.15, 1.5],
    )

    chart_path = CROSS_CHART_DIR / "2000-C收益对比折线图.png"
    if chart_path.exists():
        doc.add_paragraph()
        doc.add_picture(str(chart_path), width=Inches(6.45))
        cap = doc.add_paragraph("图 1｜2000 起点、策略 C：不同标的的累计收益路径")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cap.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(100, 116, 139)

    doc.add_heading("策略与标的的关系", level=1)
    doc.add_heading("1. 宽基正向标的：持有与顺势通常更匹配", level=2)
    doc.add_paragraph(
        "QQQ、SPY、DIA 的长期收益主要来自企业盈利增长、分红再投资和估值变化。对这类长期正向漂移资产，"
        "C 能完整捕捉上行周期；A 牺牲部分反弹初段，换取在短期走弱时离场。A 是否优于 C，通常取决于起点是否临近"
        "高波动下跌阶段，以及后续行情是趋势性下跌还是频繁震荡。"
    )
    doc.add_heading("2. 三倍做多：顺势过滤更有价值，但路径依赖被放大", level=2)
    doc.add_paragraph(
        "TQQQ、UPRO、UDOW 的每日杠杆会同时放大收益、回撤与波动损耗。C 在持续牛市中可能取得极高终值，"
        "但大幅回撤会显著降低复利基数；A 的短均线过滤往往能削弱部分急跌暴露。另一方面，5 日均线很敏感，"
        "震荡市中的反复进出会造成踏空与“来回打脸”，即使本模型没有计入成本，也会体现在收益路径中。"
    )
    doc.add_heading("3. 三倍反向：长期持有与市场正向漂移相冲突", level=2)
    doc.add_paragraph(
        "SQQQ、SPXU、SDOW 的目标是底层每日反向 3 倍，而不是长期累计反向 3 倍。底层长期上涨时，"
        "反向产品同时承受方向性逆风和每日复利衰减，因此 C 往往趋近于损失大部分本金。A 与 B 的结果取决于"
        "它们是否能让反向产品只在下跌窗口中暴露；这也说明此类产品更接近短期战术工具，而非长期核心持仓。"
    )
    doc.add_heading("4. 单一公司：趋势规则不能替代基本面", level=2)
    doc.add_paragraph(
        "KO 与 BRK.B 的波动结构、分红政策和行业暴露不同。KO 的总收益包含分红再投资，BRK.B 不派息，"
        "但通过资本配置积累价值。5 日均线只读取价格路径，无法识别估值、盈利质量或管理层变化；"
        "因此策略比较反映的是路径特征，而不是对公司基本面的判断。"
    )

    doc.add_heading("起点与市场环境", level=1)
    year_summary = (
        summary.groupby(["StartYear", "Strategy"])
        .agg(MedianCAGR=("CAGR", "median"), MedianDrawdown=("MaxDrawdown", "median"))
        .reset_index()
    )
    year_rows = []
    for year in START_YEARS:
        subdf = year_summary[year_summary["StartYear"] == year]
        best = subdf.loc[subdf["MedianCAGR"].idxmax()]
        safest = subdf.loc[subdf["MedianDrawdown"].idxmax()]
        year_rows.append([str(year), str(best["Strategy"]), f"{best['MedianCAGR']:.1%}", str(safest["Strategy"])])
    add_table(doc, ["起始年份", "中位 CAGR 最佳", "中位 CAGR", "中位回撤最小"], year_rows, [1.1, 1.5, 1.3, 1.5])
    doc.add_paragraph(
        "2000、2007、2008、2022 等起点靠近估值收缩或下跌阶段，择时规则更可能体现防御价值；"
        "2010、2018、2023 等起点之后若出现较长上行段，持续持有的机会成本更低。"
        "因此，单一起点的胜负不应被视为策略的稳定结论，多起点结果比某一条曲线更重要。"
    )

    chart_path = CROSS_CHART_DIR / "2022-A收益对比折线图.png"
    if chart_path.exists():
        doc.add_picture(str(chart_path), width=Inches(6.45))
        cap = doc.add_paragraph("图 2｜2022 起点、策略 A：不同标的的累计收益路径")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cap.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(100, 116, 139)

    doc.add_heading("风险、偏差与使用建议", level=1)
    caveats = [
        "模拟的 3 倍产品只复制每日倍数，不包含真实基金费率、融资成本、掉期价差、跟踪误差和流动性影响，因此长期终值可能明显偏乐观。",
        "现金收益设为 0。在高利率年份，空仓现金若能获得利息，A/B 的结果会更高。",
        "不计交易成本对 5 日均线策略尤其有利；真实执行中，频繁切换会产生摩擦。",
        "总收益序列适合比较财富增长，但实际成交使用市场价格；分红再投资、税务与成交时点会造成差异。",
        "本研究只比较一个均线长度与三条规则，没有做参数搜索，因此避免了部分过拟合，但也不代表 5 日是最优周期。",
    ]
    for item in caveats:
        doc.add_paragraph("• " + item)
    doc.add_paragraph(
        "更稳妥的使用方式是：先按资产家族看结构性规律，再看多个起点的一致性，最后才查看单个文件中的具体终值。"
        "如果用于真实决策，应加入手续费、滑点、现金利息、真实杠杆 ETF 费用，并做样本外检验。"
    )

    doc.add_heading("数据来源与复核", level=1)
    doc.add_paragraph(
        "基础标的使用 Total Real Returns 提供的逐日、分红再投资总收益序列："
    )
    for symbol, source in BASE_SOURCES.items():
        doc.add_paragraph(f"• {symbol}: {source}")
    doc.add_paragraph(
        "每个单标的 Excel 的“数据与说明”工作表记录了同一来源、回测口径和该文件的具体起止日期。"
        "历史数据 CSV、月度结果、汇总表和所有图表均由同一批解析后的序列生成。"
    )

    output = REPORT_DIR / "策略回测分析报告.docx"
    doc.save(output)
    return output


def main() -> None:
    ensure_dirs()
    base_frames = {}
    for symbol in ["QQQ", "DIA", "SPY", "KO", "BRK.B"]:
        base_frames[symbol] = parse_total_return_html(RAW_DIR / f"{symbol}_totalreturns.html")

    all_frames = dict(base_frames)
    for symbol, (underlying, leverage) in LEVERAGED.items():
        all_frames[symbol] = build_leveraged(base_frames[underlying], leverage)

    for symbol in TICKERS:
        if symbol in LEVERAGED:
            underlying, leverage = LEVERAGED[symbol]
            source = f"基于 {underlying} 总收益序列按每日 {leverage:+.0f} 倍模拟；底层来源：{BASE_SOURCES[underlying]}"
        else:
            underlying, leverage = symbol, 1.0
            source = BASE_SOURCES[symbol]
        save_history(symbol, all_frames[symbol], source, underlying, leverage)

    results = []
    summary_rows = []
    monthly_lookup: dict[tuple[str, int], pd.DataFrame] = {}
    for ticker in TICKERS:
        for start_year in START_YEARS:
            monthly, metrics, daily = run_backtest(all_frames[ticker], start_year)
            monthly_lookup[(ticker, start_year)] = monthly
            record = {
                "ticker": ticker,
                "start_year": start_year,
                "start_date": monthly.iloc[0]["Date"].strftime("%Y-%m-%d"),
                "end_date": all_frames[ticker].iloc[-1]["Date"].strftime("%Y-%m-%d"),
                "monthly": [
                    {
                        "date": row.Date.strftime("%Y-%m-%d"),
                        "A": float(row.A),
                        "B": float(row.B),
                        "C": float(row.C),
                    }
                    for row in monthly.itertuples(index=False)
                ],
                "metrics": {strategy: metrics[strategy].to_dict() for strategy in STRATEGIES},
            }
            results.append(record)
            for strategy in STRATEGIES:
                metric = metrics[strategy]
                summary_rows.append(
                    {
                        "Ticker": ticker,
                        "StartYear": start_year,
                        "Strategy": strategy,
                        "StartDate": record["start_date"],
                        "EndDate": record["end_date"],
                        "FinalReturn": metric.final_return,
                        "CAGR": metric.cagr,
                        "MaxDrawdown": metric.max_drawdown,
                        "Volatility": metric.volatility,
                        "Trades": metric.trades,
                        "InvestedPct": metric.invested_pct,
                    }
                )

            single_series = {strategy: monthly[strategy].tolist() for strategy in STRATEGIES}
            plot_lines(
                monthly["Date"].tolist(),
                single_series,
                COLORS,
                f"{ticker}｜{start_year} 起点｜三种策略累计收益",
                f"首个交易日 {record['start_date']}；数据截至 {record['end_date']}；月初观测",
                SINGLE_CHART_DIR / f"{ticker}-{start_year}—三种策略收益对比折线图.png",
            )

    for start_year in START_YEARS:
        for strategy in STRATEGIES:
            frames = []
            for ticker in TICKERS:
                monthly = monthly_lookup[(ticker, start_year)][["Date", strategy]].rename(columns={strategy: ticker})
                frames.append(monthly.set_index("Date"))
            aligned = pd.concat(frames, axis=1).sort_index().ffill()
            aligned = aligned.dropna(how="all")
            series = {ticker: aligned[ticker].tolist() for ticker in TICKERS}
            colors = {ticker: CROSS_COLORS[i] for i, ticker in enumerate(TICKERS)}
            plot_lines(
                aligned.index.tolist(),
                series,
                colors,
                f"{start_year} 起点｜策略 {strategy}｜11 个标的累计收益对比",
                "月初观测；三倍多空标的为基于底层总收益的每日倍数模拟",
                CROSS_CHART_DIR / f"{start_year}-{strategy}收益对比折线图.png",
            )

    summary = pd.DataFrame(summary_rows)
    summary.to_csv(OUTPUT_DIR / "策略回测汇总.csv", index=False, encoding="utf-8-sig")
    metadata = {
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "data_end": min(frame.iloc[-1]["Date"] for frame in all_frames.values()).strftime("%Y-%m-%d"),
        "start_years": START_YEARS,
        "tickers": TICKERS,
        "strategies": {
            "A": "起点满仓；前一交易日价格高于5日均线则下一交易日持有，低于则空仓",
            "B": "起点满仓；前一交易日价格低于5日均线则下一交易日持有，高于则空仓",
            "C": "起点满仓并持续持有",
        },
        "assumptions": [
            "1月1日休市时使用其后第一个交易日",
            "信号于收盘确认，下一交易日生效",
            "现金收益为0，不计交易费、税费与滑点",
            "三倍序列按底层每日总收益乘以+3或-3模拟，不计费用、融资和跟踪误差",
            "价格等于5日均线时维持上一日仓位",
        ],
        "sources": BASE_SOURCES,
    }
    (RESULT_DIR / "backtests.json").write_text(
        json.dumps({"metadata": metadata, "results": results}, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    report_path = build_report(summary, metadata["data_end"], results)

    manifest = {
        "history_csv": len(list(DATA_DIR.glob("*.csv"))),
        "single_charts": len(list(SINGLE_CHART_DIR.glob("*.png"))),
        "cross_charts": len(list(CROSS_CHART_DIR.glob("*.png"))),
        "results": len(results),
        "report": str(report_path),
        "data_ranges": {
            symbol: {
                "start": frame.iloc[0]["Date"].strftime("%Y-%m-%d"),
                "end": frame.iloc[-1]["Date"].strftime("%Y-%m-%d"),
                "rows": int(len(frame)),
            }
            for symbol, frame in all_frames.items()
        },
    }
    (RESULT_DIR / "manifest.json").write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(manifest, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
